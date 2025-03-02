import xml.etree.ElementTree as ET
import os
from docx import Document
from docx.shared import RGBColor
import json

def create_word_with_citation_markers(xml_path, references_data=None, pmcid=None):
    """创建带有精确匹配引用标记的Word文档，使用传入的PMCID创建唯一标识符"""
    tree = ET.parse(xml_path)
    root = tree.getroot()
    doc = Document()
    
    # 如果没有提供PMCID，使用默认值
    if pmcid is None:
        pmcid = "UNKNOWN"
        print("警告: 未提供PMCID，使用默认值'UNKNOWN'")
    
    # 如果没有提供引用数据，尝试从之前生成的JSON加载
    if references_data is None:
        json_path = os.path.splitext(xml_path)[0] + "_refs.json"
        if os.path.exists(json_path):
            with open(json_path, 'r', encoding='utf-8') as f:
                references_data = json.load(f)
        else:
            print("未找到引用数据，将无法正确标记引用")
            references_data = []
    
    # 创建引用ID到索引的映射
    citations = {ref['id']: i+1 for i, ref in enumerate(references_data)}
    
    # 提取文章元数据
    article_meta = root.find('.//article-meta')
    
    # 提取标题
    title_elem = article_meta.find('.//article-title')
    title = title_elem.text if title_elem is not None else "Untitled"
    doc.add_heading(title, level=0)
    
    # 提取作者
    authors = []
    for contrib in article_meta.findall('.//contrib[@contrib-type="author"]'):
        surname = contrib.find('.//surname')
        given_names = contrib.find('.//given-names')
        if surname is not None and given_names is not None:
            authors.append(f"{given_names.text} {surname.text}")
    
    doc.add_paragraph(", ".join(authors))
    
    # 提取摘要
    abstract = ""
    abstract_elem = article_meta.find('.//abstract')
    if abstract_elem is not None:
        doc.add_heading("Abstract", level=1)
        abstract_para = doc.add_paragraph()
        
        for p in abstract_elem.findall('.//p'):
            if p.text:
                abstract_para.add_run(p.text + " ")
    
    # 提取正文
    body = root.find('.//body')
    if body is not None:
        # 处理每个段落
        for section in body.findall('.//sec'):
            # 提取段落标题
            title = section.find('.//title')
            if title is not None and title.text:
                doc.add_heading(title.text, level=1)
            
            # 处理段落内容
            for p in section.findall('.//p'):
                para = doc.add_paragraph()
                
                # 使用新的处理函数处理段落内容
                process_paragraph_with_ranges(p, para, citations, pmcid)
    
    # 保存文档
    docx_path = os.path.splitext(xml_path)[0] + ".docx"
    doc.save(docx_path)
    
    # 保存引用数据以便后续使用
    if references_data and not os.path.exists(os.path.splitext(xml_path)[0] + "_refs.json"):
        with open(os.path.splitext(xml_path)[0] + "_refs.json", 'w', encoding='utf-8') as f:
            json.dump(references_data, f, ensure_ascii=False, indent=2)
    
    print(f"Word文档已创建: {docx_path}")
    return docx_path

def process_paragraph_with_ranges(element, paragraph, citations, pmcid):
    """处理段落内容，特别处理引用范围，使用PMCID创建唯一标识符"""
    # 将元素转换为字符串以便于分析
    element_str = ET.tostring(element, encoding='unicode')
    
    # 创建一个新的解析函数来处理元素内容
    def parse_element(elem, current_para):
        # 处理元素前的文本
        if elem.text:
            current_para.add_run(elem.text)
        
        # 用于跟踪是否在处理引用范围
        in_range = False
        range_start = None
        
        # 处理所有子元素
        i = 0
        while i < len(elem):
            child = elem[i]
            
            # 检查是否是引用元素
            if child.tag == 'xref' and child.get('ref-type') == 'bibr':
                ref_id = child.get('rid')
                
                # 创建唯一标识符
                unique_id = f"{pmcid}_{ref_id}"
                
                # 检查是否可能是范围的开始
                if i + 1 < len(elem):
                    # 获取当前元素和下一个元素之间的文本
                    between_text = child.tail if child.tail else ""
                    
                    # 检查下一个元素是否也是引用
                    next_child = elem[i + 1]
                    if (next_child.tag == 'xref' and 
                        next_child.get('ref-type') == 'bibr' and
                        ('–' in between_text or '-' in between_text or '&#x2013;' in between_text)):
                        
                        # 这是一个引用范围的开始
                        in_range = True
                        range_start = ref_id
                        
                        # 获取范围结束
                        range_end = next_child.get('rid')
                        
                        # 提取数字部分
                        if range_start.startswith('R') and range_end.startswith('R'):
                            start_num = int(range_start[1:])
                            end_num = int(range_end[1:])
                            
                            # 生成范围内所有引用的标记，使用PMCID
                            citation_markers = ""
                            for num in range(start_num, end_num + 1):
                                current_id = f"R{num}"
                                unique_current_id = f"{pmcid}_{current_id}"
                                if current_id in citations:
                                    citation_markers += f"{{^{unique_current_id}$}}"
                            
                            # 添加所有引用标记
                            if citation_markers:
                                run = current_para.add_run(citation_markers)
                                run.bold = True
                                run.font.color.rgb = RGBColor(0, 0, 255)
                            
                            # 处理范围结束后的文本
                            if next_child.tail:
                                current_para.add_run(next_child.tail)
                            
                            # 跳过下一个元素，因为已经处理了
                            i += 2
                            continue
                
                # 如果不是范围或范围已结束
                if not in_range:
                    # 单个引用，使用PMCID
                    if ref_id in citations:
                        citation_marker = f"{{^{unique_id}$}}"
                        run = current_para.add_run(citation_marker)
                        run.bold = True
                        run.font.color.rgb = RGBColor(0, 0, 255)
                    
                    # 处理引用后的文本
                    if child.tail:
                        current_para.add_run(child.tail)
            else:
                # 递归处理其他元素
                sub_para = current_para
                parse_element(child, sub_para)
                
                # 处理元素后的文本
                if child.tail:
                    current_para.add_run(child.tail)
            
            # 移动到下一个元素
            i += 1
    
    # 开始解析，传入PMCID
    parse_element(element, paragraph)

if __name__ == "__main__":
    xml_file = "repodir/PMC7096285.xml"  # 替换为您的XML文件路径
    
    # 首先提取引用数据
    from generate_ris import extract_references_to_ris
    _, references, pmcid = extract_references_to_ris(xml_file)
    
    # 然后创建Word文档
    docx_file = create_word_with_citation_markers(xml_file, references, pmcid)
    print(f"Word文档已生成: {docx_file}")